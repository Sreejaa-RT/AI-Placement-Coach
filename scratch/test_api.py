import os
import requests
import docx

# Base API URL
API_URL = "http://localhost:8000/api/v1/resume/audit"
SCRATCH_DIR = os.path.dirname(__file__)

def generate_mock_documents():
    """
    Generates required mock files for API testing under scratch/ directory.
    """
    print("[Test Prep] Generating mock test files using python-docx...")
    
    # Ensure scratch directory exists
    os.makedirs(SCRATCH_DIR, exist_ok=True)
    
    # 1. Strong Match Resume
    doc1 = docx.Document()
    doc1.add_heading("John Doe - Senior Data Scientist", 0)
    doc1.add_paragraph(
        "Professional Summary:\n"
        "Over 5 years of industry experience designing and deploying predictive machine learning models.\n"
        "Technical Skills:\n"
        "- Languages: Python, SQL, C++\n"
        "- Frameworks: Scikit-learn, Pandas, NumPy, TensorFlow, PyTorch\n"
        "- Tools: Git, Docker, AWS Cloud infrastructure\n"
        "Experience:\n"
        "- Led development of neural network models using Deep Learning to automate product classifications."
    )
    doc1_path = os.path.join(SCRATCH_DIR, "strong_ds.docx")
    doc1.save(doc1_path)
    
    # 2. Moderate Match Resume
    doc2 = docx.Document()
    doc2.add_heading("Jane Smith - Software Developer", 0)
    doc2.add_paragraph(
        "Professional Summary:\n"
        "Software engineer with 3 years of experience writing web backends and scripts.\n"
        "Technical Skills:\n"
        "- Languages: Python, JavaScript, Java, SQL\n"
        "- Tools: Git, GitHub, MySQL\n"
        "Experience:\n"
        "- Developed relational database query schedules and integrated API endpoints in Spring Boot."
    )
    doc2_path = os.path.join(SCRATCH_DIR, "moderate_web.docx")
    doc2.save(doc2_path)
    
    # 3. Poor Match Resume
    doc3 = docx.Document()
    doc3.add_heading("Robert Johnson - Accountant & Finance Officer", 0)
    doc3.add_paragraph(
        "Professional Summary:\n"
        "Focused Corporate Accountant with 4 years experience managing general ledgers.\n"
        "Skills: Bookkeeping, tax filing, Accounts Payable (AP), Accounts Receivable (AR), Microsoft Excel, spreadsheet validation."
    )
    doc3_path = os.path.join(SCRATCH_DIR, "poor_acct.docx")
    doc3.save(doc3_path)
    
    # 4. Plain Text File (unsupported format)
    txt_path = os.path.join(SCRATCH_DIR, "invalid_format.txt")
    with open(txt_path, "w") as f:
        f.write("This is a plain text resume file which should be rejected by the backend parser.")
        
    # 5. Empty File
    empty_path = os.path.join(SCRATCH_DIR, "empty.docx")
    doc_empty = docx.Document()
    # Save a docx with no paragraph content (will result in empty text)
    doc_empty.save(empty_path)

    print("[Test Prep] Mock files successfully created:")
    print(f"  - {doc1_path}")
    print(f"  - {doc2_path}")
    print(f"  - {doc3_path}")
    print(f"  - {txt_path}")
    print(f"  - {empty_path}")
    print("-" * 50)
    
    return doc1_path, doc2_path, doc3_path, txt_path, empty_path

def run_api_tests():
    # Pre-generate mock files
    strong_p, mod_p, poor_p, txt_p, empty_p = generate_mock_documents()
    
    # Target Job Role ID for tests: 5 (Data Scientist)
    target_role_id = 5
    
    # ----------------------------------------------------
    # TEST 1: Strong Data Scientist resume (Expect 200)
    # ----------------------------------------------------
    print("\n[TEST 1] Uploading Strong Data Scientist Resume...")
    with open(strong_p, "rb") as f:
        files = {"resume": (os.path.basename(strong_p), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"job_role_id": target_role_id}
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    assert res.status_code == 200, f"Expected 200, got {res.status_code}"
    json_data = res.json()
    print(f"Resume Fit Score  : {json_data['resume_fit_score']}/100")
    print(f"Fit Category      : {json_data['fit_category']}")
    print(f"Predicted Label   : {json_data['predicted_label']}")
    print(f"Skill Match Ratio : {json_data['skill_match_ratio']:.4f}")
    print(f"Matched Skills    : {json_data['matched_skills']}")
    print(f"Missing Skills    : {json_data['missing_skills']}")
    print(f"Strengths         : {json_data['strengths']}")
    print(f"Suggestions       : {json_data['suggestions']}")
    prob_sum = json_data['good_fit_probability'] + json_data['potential_fit_probability'] + json_data['no_fit_probability']
    print(f"Probability Sum   : {prob_sum:.4f}")
    assert 0.99 <= prob_sum <= 1.01, f"Expected probabilities to sum to ~1.0, got {prob_sum}"
    
    # ----------------------------------------------------
    # TEST 2: Moderate Web Developer resume (Expect 200)
    # ----------------------------------------------------
    print("\n[TEST 2] Uploading Moderate Web Developer Resume...")
    with open(mod_p, "rb") as f:
        files = {"resume": (os.path.basename(mod_p), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"job_role_id": target_role_id}
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    assert res.status_code == 200, f"Expected 200, got {res.status_code}"
    json_data = res.json()
    print(f"Resume Fit Score  : {json_data['resume_fit_score']}/100")
    print(f"Fit Category      : {json_data['fit_category']}")
    print(f"Predicted Label   : {json_data['predicted_label']}")
    print(f"Skill Match Ratio : {json_data['skill_match_ratio']:.4f}")
    print(f"Matched Skills    : {json_data['matched_skills']}")
    print(f"Missing Skills    : {json_data['missing_skills']}")
    
    # ----------------------------------------------------
    # TEST 3: Poor Accountant resume (Expect 200)
    # ----------------------------------------------------
    print("\n[TEST 3] Uploading Poor Accountant Resume...")
    with open(poor_p, "rb") as f:
        files = {"resume": (os.path.basename(poor_p), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"job_role_id": target_role_id}
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    assert res.status_code == 200, f"Expected 200, got {res.status_code}"
    json_data = res.json()
    print(f"Resume Fit Score  : {json_data['resume_fit_score']}/100")
    print(f"Fit Category      : {json_data['fit_category']}")
    print(f"Predicted Label   : {json_data['predicted_label']}")
    print(f"Skill Match Ratio : {json_data['skill_match_ratio']:.4f}")
    print(f"Matched Skills    : {json_data['matched_skills']}")
    print(f"Missing Skills    : {json_data['missing_skills']}")
    
    # ----------------------------------------------------
    # TEST 4: TXT Upload Rejection (Expect 400)
    # ----------------------------------------------------
    print("\n[TEST 4] Uploading TXT file (Unsupported Extension)...")
    with open(txt_p, "rb") as f:
        files = {"resume": (os.path.basename(txt_p), f, "text/plain")}
        data = {"job_role_id": target_role_id}
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    print(f"Error Msg  : {res.json()['detail']}")
    assert res.status_code == 400, f"Expected 400, got {res.status_code}"
    
    # ----------------------------------------------------
    # TEST 5: Empty DOCX Rejection (Expect 400)
    # ----------------------------------------------------
    print("\n[TEST 5] Uploading Empty/No-Text DOCX file...")
    with open(empty_p, "rb") as f:
        files = {"resume": (os.path.basename(empty_p), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"job_role_id": target_role_id}
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    print(f"Error Msg  : {res.json()['detail']}")
    assert res.status_code == 400, f"Expected 400, got {res.status_code}"
    
    # ----------------------------------------------------
    # TEST 6: Invalid job_role_id Rejection (Expect 400)
    # ----------------------------------------------------
    print("\n[TEST 6] Requesting with Invalid job_role_id...")
    with open(strong_p, "rb") as f:
        files = {"resume": (os.path.basename(strong_p), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"job_role_id": 99} # Invalid role index
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    print(f"Error Msg  : {res.json()['detail']}")
    assert res.status_code == 400, f"Expected 400, got {res.status_code}"
    
    # ----------------------------------------------------
    # TEST 7: Custom Role without Job Description (Expect 400)
    # ----------------------------------------------------
    print("\n[TEST 7] Requesting Custom Role (id=13) without job description...")
    with open(strong_p, "rb") as f:
        files = {"resume": (os.path.basename(strong_p), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"job_role_id": 13} # Custom Role
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    print(f"Error Msg  : {res.json()['detail']}")
    assert res.status_code == 400, f"Expected 400, got {res.status_code}"
    
    # ----------------------------------------------------
    # TEST 8: Custom Role WITH Job Description (Expect 200)
    # ----------------------------------------------------
    print("\n[TEST 8] Requesting Custom Role (id=13) WITH valid job description...")
    with open(strong_p, "rb") as f:
        files = {"resume": (os.path.basename(strong_p), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {
            "job_role_id": 13,
            "custom_job_description": "We need a Software developer skilled in C++ and Docker with Git workflow."
        }
        res = requests.post(API_URL, files=files, data=data)
        
    print(f"Status Code: {res.status_code}")
    assert res.status_code == 200, f"Expected 200, got {res.status_code}"
    json_data = res.json()
    print(f"Resume Fit Score  : {json_data['resume_fit_score']}/100")
    print(f"Fit Category      : {json_data['fit_category']}")
    print(f"Matched Skills    : {json_data['matched_skills']}")
    
    print("\n==================================================")
    print("ALL API VALIDATION TESTS PASSED SUCCESSFULLY!")
    print("==================================================")

if __name__ == "__main__":
    run_api_tests()
